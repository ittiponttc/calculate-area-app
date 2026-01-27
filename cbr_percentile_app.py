import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from scipy import interpolate
import json
import base64
import io

# Check if python-docx is available
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(
    page_title="CBR Percentile Analysis",
    page_icon="📊",
    layout="wide"
)

st.title("📊 การวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์")
st.markdown("### Subgrade CBR Analysis Tool")
st.markdown("---")

# Sample data (CBR values only)
sample_cbr = [14.8, 14.37, 5.31, 17.37, 5.48, 18.46, 4.85, 6.23,
              5.02, 10.78, 10.52, 14, 15.5, 8.7, 12.93, 8.19,
              8.1, 15.56, 16.88, 20.75, 20.3, 8, 7.84, 7.48,
              23.55, 8.92, 13.3, 13.5, 13.86, 7.18, 6.95, 5.8,
              6, 11.18, 9.69, 7.48]

# Sidebar for file upload
with st.sidebar:
    st.header("📁 อัปโหลดข้อมูล")
    
    # Upload JSON for settings
    st.markdown("#### 📂 โหลดการตั้งค่า")
    uploaded_json = st.file_uploader(
        "โหลดข้อมูลจากไฟล์ JSON",
        type=['json'],
        help="โหลดค่า Percentile และข้อมูล CBR จากไฟล์ JSON"
    )
    
    if uploaded_json is not None:
        try:
            loaded_data = json.load(uploaded_json)
            
            # ตรวจสอบว่าเป็นไฟล์ใหม่
            file_id = f"{uploaded_json.name}_{uploaded_json.size}"
            if st.session_state.get('last_uploaded_json') != file_id:
                st.session_state['last_uploaded_json'] = file_id
                
                # อัพเดท session_state
                if 'target_percentile' in loaded_data:
                    st.session_state['input_percentile'] = float(loaded_data['target_percentile'])
                if 'cbr_values' in loaded_data:
                    st.session_state['loaded_cbr_values'] = loaded_data['cbr_values']
                if 'use_sample' in loaded_data:
                    st.session_state['input_use_sample'] = loaded_data['use_sample']
                
                st.success("✅ โหลดการตั้งค่าสำเร็จ!")
                st.rerun()
                
        except Exception as e:
            st.error(f"❌ ไม่สามารถอ่านไฟล์ JSON ได้: {e}")
    
    st.markdown("---")
    
    # Upload Excel for CBR data
    st.markdown("#### 📊 อัปโหลดข้อมูล CBR")
    uploaded_file = st.file_uploader(
        "เลือกไฟล์ Excel (.xlsx)",
        type=['xlsx'],
        help="ไฟล์ควรมีคอลัมน์ CBR(%) เพียงคอลัมน์เดียว"
    )
    
    st.markdown("---")
    st.markdown("### 📋 รูปแบบข้อมูลที่ต้องการ")
    st.markdown("""
    | CBR(%) |
    |--------|
    | 14.8   |
    | 14.37  |
    | 5.31   |
    | ...    |
    """)
    st.info("ระบบจะคำนวณ Percentile ให้อัตโนมัติ")

# Process uploaded Excel file
if uploaded_file is not None:
    try:
        # Read Excel file
        df = pd.read_excel(uploaded_file)
        
        # Try to identify CBR column
        cbr_col = None
        
        for col in df.columns:
            col_lower = str(col).lower()
            if 'cbr' in col_lower:
                cbr_col = col
                break
        
        # If not found, use first column
        if cbr_col is None:
            cbr_col = df.columns[0]
        
        # Get CBR values
        cbr_values = pd.to_numeric(df[cbr_col], errors='coerce').dropna().tolist()
        
        st.success(f"✅ อ่านข้อมูลสำเร็จ: {len(cbr_values)} ตัวอย่าง")
        
    except Exception as e:
        st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
        st.info("กรุณาตรวจสอบรูปแบบไฟล์ Excel")
        cbr_values = None

elif 'loaded_cbr_values' in st.session_state and st.session_state['loaded_cbr_values']:
    # Use CBR values from loaded JSON
    cbr_values = st.session_state['loaded_cbr_values']
    st.info(f"📌 ใช้ข้อมูลจากไฟล์ JSON: {len(cbr_values)} ตัวอย่าง")

else:
    st.info("📌 กรุณาอัปโหลดไฟล์ Excel หรือใช้ข้อมูลตัวอย่าง")
    
    default_use_sample = st.session_state.get('input_use_sample', True)
    use_sample = st.checkbox(
        "ใช้ข้อมูลตัวอย่าง", 
        value=default_use_sample,
        key="input_use_sample"
    )
    
    if use_sample:
        cbr_values = sample_cbr
    else:
        cbr_values = None

if cbr_values is not None and len(cbr_values) > 0:
    
    # Sort CBR values
    cbr_sorted = np.sort(cbr_values)
    n = len(cbr_sorted)
    
    # Calculate cumulative percentile (percentage of values <= each CBR)
    cumulative_percentile = (np.arange(1, n + 1) / n) * 100
    
    # Create dataframe for display
    df_sorted = pd.DataFrame({
        'CBR': cbr_sorted,
        'Cumulative_Percentile': cumulative_percentile
    })
    
    # Create interpolation function
    f_interp = interpolate.interp1d(
        cumulative_percentile, 
        cbr_sorted,
        kind='linear',
        fill_value='extrapolate'
    )
    
    # Input percentile at the top
    st.markdown("### 🎯 กำหนดค่า Percentile")
    
    default_percentile = st.session_state.get('input_percentile', 90.0)
    target_percentile = st.number_input(
        "Percentile ที่ต้องการ (%)",
        min_value=0.0,
        max_value=100.0,
        value=default_percentile,
        step=1.0,
        help="ใส่ค่า Percentile ที่ต้องการหาค่า CBR",
        key="input_percentile"
    )
    
    # Calculate CBR at target percentile
    design_percentile = 100 - target_percentile
    
    if design_percentile >= cumulative_percentile.min() and \
       design_percentile <= cumulative_percentile.max():
        cbr_at_percentile = float(f_interp(design_percentile))
    else:
        cbr_at_percentile = float(f_interp(np.clip(design_percentile, 
                                                    cumulative_percentile.min(),
                                                    cumulative_percentile.max())))
    
    st.markdown("---")
    
    # Graph section - full width
    st.markdown("### 📈 กราฟ Percentile vs CBR")
    
    # Create figure
    fig = go.Figure()
    
    # Calculate axis ranges
    x_max = max(cbr_sorted) * 1.1
    y_max = 100
    
    # Add main curve
    fig.add_trace(go.Scatter(
        x=cbr_sorted,
        y=100 - cumulative_percentile,  # Convert to "% >= value"
        mode='lines+markers',
        name='CBR Distribution',
        line=dict(color='blue', width=2),
        marker=dict(size=6, symbol='x', color='black')
    ))
    
    # Add horizontal red dashed line at target percentile
    fig.add_trace(go.Scatter(
        x=[0, cbr_at_percentile],
        y=[target_percentile, target_percentile],
        mode='lines',
        name=f'Percentile {target_percentile}%',
        line=dict(color='red', width=2, dash='dash')
    ))
    
    # Add vertical red dashed line at CBR value
    fig.add_trace(go.Scatter(
        x=[cbr_at_percentile, cbr_at_percentile],
        y=[0, target_percentile],
        mode='lines',
        name=f'CBR = {cbr_at_percentile:.2f}%',
        line=dict(color='red', width=2, dash='dash')
    ))
    
    # Add annotation for CBR value
    fig.add_annotation(
        x=cbr_at_percentile,
        y=0,
        text=f"<b>{cbr_at_percentile:.2f}</b>",
        showarrow=True,
        arrowhead=2,
        arrowsize=1,
        arrowwidth=2,
        arrowcolor='red',
        ax=0,
        ay=40,
        font=dict(size=16, color='red')
    )
    
    # Border line width (consistent for all 4 sides)
    border_width = 4
    
    # Update layout - remove axis lines, we'll draw border using shapes
    fig.update_layout(
        xaxis_title="CBR (%)",
        yaxis_title="Percentile (%)",
        xaxis=dict(
            range=[0, x_max],
            gridcolor='lightgray',
            showgrid=True,
            showline=False,  # Disable built-in axis line
            zeroline=False,
            ticks='outside',
            tickwidth=1,
            tickcolor='black',
            ticklen=5,
        ),
        yaxis=dict(
            range=[0, y_max],
            gridcolor='lightgray',
            showgrid=True,
            showline=False,  # Disable built-in axis line
            zeroline=False,
            ticks='outside',
            tickwidth=1,
            tickcolor='black',
            ticklen=5,
        ),
        plot_bgcolor='white',
        width=600,
        height=600,
        showlegend=True,
        legend=dict(
            yanchor="top",
            y=0.99,
            xanchor="right",
            x=0.99,
            bgcolor='rgba(255,255,255,0.8)',
            bordercolor='black',
            borderwidth=1
        ),
        title=dict(
            text=f"ค่าร้อยละ CBR ที่เปอร์เซ็นต์ไทล์ ร้อยละ {target_percentile:.0f}",
            x=0.5,
            xanchor='center'
        ),
        margin=dict(l=70, r=70, t=70, b=70)
    )
    
    # Draw complete border using a rectangle shape (ensures all 4 corners connect)
    fig.add_shape(
        type="rect",
        x0=0, y0=0,
        x1=x_max, y1=y_max,
        line=dict(color="black", width=border_width),
        xref="x", yref="y"
    )
    
    # Center the chart
    col_left, col_chart, col_right = st.columns([1, 2, 1])
    with col_chart:
        st.plotly_chart(fig, use_container_width=False)
    
    # Results section - below the graph
    st.markdown("---")
    
    col_result, col_stat = st.columns(2)
    
    with col_result:
        st.markdown("### 📊 ผลการวิเคราะห์")
        st.metric(
            label=f"CBR ที่ Percentile {target_percentile}%",
            value=f"{cbr_at_percentile:.2f} %"
        )
    
    with col_stat:
        st.markdown("### 📋 สถิติข้อมูล CBR")
        st.write(f"**จำนวนตัวอย่าง:** {n}")
        st.write(f"**ค่าต่ำสุด:** {np.min(cbr_values):.2f} %")
        st.write(f"**ค่าสูงสุด:** {np.max(cbr_values):.2f} %")
        st.write(f"**ค่าเฉลี่ย:** {np.mean(cbr_values):.2f} %")
        st.write(f"**ส่วนเบี่ยงเบนมาตรฐาน:** {np.std(cbr_values):.2f} %")
    
    # Export section
    st.markdown("---")
    st.markdown("### 💾 บันทึกข้อมูล")
    
    col_json, col_word = st.columns(2)
    
    with col_json:
        # Prepare export data for JSON
        export_data = {
            'target_percentile': target_percentile,
            'cbr_at_percentile': round(cbr_at_percentile, 2),
            'cbr_values': [float(v) for v in cbr_values],
            'statistics': {
                'n_samples': n,
                'min': round(float(np.min(cbr_values)), 2),
                'max': round(float(np.max(cbr_values)), 2),
                'mean': round(float(np.mean(cbr_values)), 2),
                'std': round(float(np.std(cbr_values)), 2)
            },
            'use_sample': st.session_state.get('input_use_sample', True)
        }
        
        json_str = json.dumps(export_data, ensure_ascii=False, indent=2)
        
        st.download_button(
            label="📥 Download JSON",
            data=json_str,
            file_name="cbr_percentile_data.json",
            mime="application/json",
            help="บันทึกข้อมูลและการตั้งค่าเป็นไฟล์ JSON"
        )
    
    with col_word:
        # Generate Word document using python-docx
        if DOCX_AVAILABLE:
            if st.button("📄 สร้างรายงาน Word", help="สร้างรายงานผลการวิเคราะห์เป็นไฟล์ Word"):
                try:
                    # Create Word document
                    doc = DocxDocument()
                    
                    # Set Thai font style
                    style = doc.styles['Normal']
                    style.font.name = 'TH SarabunPSK'
                    style.font.size = Pt(16)
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH SarabunPSK')
                    
                    # Title
                    title = doc.add_heading('', level=1)
                    title_run = title.add_run('รายงานผลการวิเคราะห์ค่า CBR ที่เปอร์เซ็นต์ไทล์')
                    title_run.font.name = 'TH SarabunPSK'
                    title_run.font.size = Pt(20)
                    title_run.font.bold = True
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Results section
                    h2 = doc.add_heading('', level=2)
                    h2_run = h2.add_run('ผลการวิเคราะห์')
                    h2_run.font.name = 'TH SarabunPSK'
                    h2_run.font.size = Pt(18)
                    h2_run.font.bold = True
                    
                    p1 = doc.add_paragraph()
                    p1_run1 = p1.add_run('Percentile ที่กำหนด: ')
                    p1_run1.font.name = 'TH SarabunPSK'
                    p1_run1.font.size = Pt(16)
                    p1_run1.font.bold = True
                    p1_run2 = p1.add_run(f'{target_percentile:.0f} %')
                    p1_run2.font.name = 'TH SarabunPSK'
                    p1_run2.font.size = Pt(16)
                    
                    p2 = doc.add_paragraph()
                    p2_run1 = p2.add_run(f'ค่า CBR ที่ Percentile {target_percentile:.0f}%: ')
                    p2_run1.font.name = 'TH SarabunPSK'
                    p2_run1.font.size = Pt(16)
                    p2_run1.font.bold = True
                    p2_run2 = p2.add_run(f'{cbr_at_percentile:.2f} %')
                    p2_run2.font.name = 'TH SarabunPSK'
                    p2_run2.font.size = Pt(16)
                    p2_run2.font.bold = True
                    p2_run2.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Statistics section
                    h3 = doc.add_heading('', level=2)
                    h3_run = h3.add_run('สถิติข้อมูล CBR')
                    h3_run.font.name = 'TH SarabunPSK'
                    h3_run.font.size = Pt(18)
                    h3_run.font.bold = True
                    
                    # Create table
                    table = doc.add_table(rows=7, cols=2)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Table data
                    table_data = [
                        ('รายการ', 'ค่า'),
                        ('จำนวนตัวอย่าง', f'{n}'),
                        ('ค่าต่ำสุด', f'{np.min(cbr_values):.2f} %'),
                        ('ค่าสูงสุด', f'{np.max(cbr_values):.2f} %'),
                        ('ค่าเฉลี่ย', f'{np.mean(cbr_values):.2f} %'),
                        ('ส่วนเบี่ยงเบนมาตรฐาน', f'{np.std(cbr_values):.2f} %'),
                        (f'CBR ที่ Percentile {target_percentile:.0f}%', f'{cbr_at_percentile:.2f} %')
                    ]
                    
                    for i, (col1, col2) in enumerate(table_data):
                        row = table.rows[i]
                        cell1 = row.cells[0]
                        cell2 = row.cells[1]
                        
                        run1 = cell1.paragraphs[0].add_run(col1)
                        run1.font.name = 'TH SarabunPSK'
                        run1.font.size = Pt(14)
                        if i == 0:
                            run1.font.bold = True
                        
                        run2 = cell2.paragraphs[0].add_run(col2)
                        run2.font.name = 'TH SarabunPSK'
                        run2.font.size = Pt(14)
                        if i == 0:
                            run2.font.bold = True
                        if i == 6:  # Last row - CBR result
                            run2.font.bold = True
                            run2.font.color.rgb = RGBColor(255, 0, 0)
                    
                    # Set column widths
                    for row in table.rows:
                        row.cells[0].width = Cm(6)
                        row.cells[1].width = Cm(4)
                    
                    # Add note about chart
                    doc.add_paragraph()
                    h4 = doc.add_heading('', level=2)
                    h4_run = h4.add_run('กราฟ Percentile vs CBR')
                    h4_run.font.name = 'TH SarabunPSK'
                    h4_run.font.size = Pt(18)
                    h4_run.font.bold = True
                    
                    # Create chart using matplotlib
                    fig_mpl, ax = plt.subplots(figsize=(6, 6))
                    
                    # Plot main curve
                    y_plot = 100 - cumulative_percentile
                    ax.plot(cbr_sorted, y_plot, 'b-', linewidth=2, marker='x', 
                           markersize=6, markerfacecolor='black', markeredgecolor='black',
                           label='CBR Distribution')
                    
                    # Plot dashed lines
                    ax.plot([0, cbr_at_percentile], [target_percentile, target_percentile], 
                           'r--', linewidth=2, label=f'Percentile {target_percentile}%')
                    ax.plot([cbr_at_percentile, cbr_at_percentile], [0, target_percentile], 
                           'r--', linewidth=2, label=f'CBR = {cbr_at_percentile:.2f}%')
                    
                    # Annotation
                    ax.annotate(f'{cbr_at_percentile:.2f}', 
                               xy=(cbr_at_percentile, 0), 
                               xytext=(cbr_at_percentile, -8),
                               fontsize=12, color='red', fontweight='bold',
                               ha='center')
                    
                    ax.set_xlim(0, max(cbr_sorted) * 1.1)
                    ax.set_ylim(0, 100)
                    ax.set_xlabel('CBR (%)', fontsize=12)
                    ax.set_ylabel('Percentile (%)', fontsize=12)
                    ax.set_title(f'CBR at Percentile {target_percentile:.0f}%', fontsize=14)
                    ax.legend(loc='upper right', fontsize=10)
                    ax.grid(True, alpha=0.3)
                    
                    # Set border
                    for spine in ax.spines.values():
                        spine.set_linewidth(2)
                        spine.set_color('black')
                    
                    plt.tight_layout()
                    
                    # Save chart to buffer
                    chart_buffer = io.BytesIO()
                    fig_mpl.savefig(chart_buffer, format='png', dpi=150, 
                                   bbox_inches='tight', facecolor='white', edgecolor='none')
                    chart_buffer.seek(0)
                    plt.close(fig_mpl)
                    
                    # Add chart image to document
                    chart_para = doc.add_paragraph()
                    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    chart_run = chart_para.add_run()
                    chart_run.add_picture(chart_buffer, width=Cm(12))
                    
                    # Add caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run('รูปที่ 1 กราฟแสดงความสัมพันธ์ระหว่าง Percentile และ CBR')
                    caption_run.font.name = 'TH SarabunPSK'
                    caption_run.font.size = Pt(14)
                    caption_run.font.italic = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add CBR data table section
                    doc.add_paragraph()
                    h5 = doc.add_heading('', level=2)
                    h5_run = h5.add_run('ตารางข้อมูล CBR (เรียงตาม CBR)')
                    h5_run.font.name = 'TH SarabunPSK'
                    h5_run.font.size = Pt(18)
                    h5_run.font.bold = True
                    
                    # Create CBR data table
                    cbr_table = doc.add_table(rows=n+1, cols=3)
                    cbr_table.style = 'Table Grid'
                    cbr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Header row
                    header_row = cbr_table.rows[0]
                    headers = ['ลำดับ', 'CBR (%)', 'Percentile (%)']
                    for j, header_text in enumerate(headers):
                        cell = header_row.cells[j]
                        run = cell.paragraphs[0].add_run(header_text)
                        run.font.name = 'TH SarabunPSK'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Data rows
                    for i in range(n):
                        row = cbr_table.rows[i+1]
                        cbr_val = cbr_sorted[i]
                        pct_val = 100 - cumulative_percentile[i]
                        
                        data = [f'{i+1}', f'{cbr_val:.2f}', f'{pct_val:.2f}']
                        for j, val in enumerate(data):
                            cell = row.cells[j]
                            run = cell.paragraphs[0].add_run(val)
                            run.font.name = 'TH SarabunPSK'
                            run.font.size = Pt(14)
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Set column widths for CBR table
                    for row in cbr_table.rows:
                        row.cells[0].width = Cm(2)
                        row.cells[1].width = Cm(3)
                        row.cells[2].width = Cm(3)
                    
                    # Footer
                    doc.add_paragraph()
                    doc.add_paragraph()
                    footer1 = doc.add_paragraph()
                    footer1_run = footer1.add_run('พัฒนาโดย รศ.ดร.อิทธิพล มีผล')
                    footer1_run.font.name = 'TH SarabunPSK'
                    footer1_run.font.size = Pt(14)
                    footer1_run.font.italic = True
                    footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    footer2 = doc.add_paragraph()
                    footer2_run = footer2.add_run('ภาควิชาครุศาสตร์โยธา คณะครุศาสตร์อุตสาหกรรม มจพ.')
                    footer2_run.font.name = 'TH SarabunPSK'
                    footer2_run.font.size = Pt(14)
                    footer2_run.font.italic = True
                    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Save to buffer
                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        label="📥 Download Word",
                        data=buffer,
                        file_name="cbr_percentile_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    st.success("✅ สร้างรายงาน Word สำเร็จ!")
                    
                except Exception as e:
                    st.error(f"❌ ไม่สามารถสร้างรายงาน Word ได้: {e}")
        else:
            st.warning("⚠️ ต้องติดตั้ง python-docx เพื่อใช้งานฟีเจอร์นี้")
            st.code("pip install python-docx", language="bash")
    
    # Show data table
    st.markdown("---")
    st.markdown("### 📋 ตารางข้อมูล (เรียงตาม CBR)")
    
    # Create display table with calculated percentile
    df_display = pd.DataFrame({
        'ลำดับ': range(1, n + 1),
        'CBR (%)': cbr_sorted,
        'Percentile (%)': np.round(100 - cumulative_percentile, 2)
    })
    
    col_a, col_b = st.columns(2)
    
    with col_a:
        st.dataframe(
            df_display.head(len(df_display)//2 + 1),
            use_container_width=True,
            hide_index=True
        )
    
    with col_b:
        st.dataframe(
            df_display.tail(len(df_display)//2),
            use_container_width=True,
            hide_index=True
        )

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray;'>
    <p>สำหรับการวิเคราะห์ค่า CBR ดินฐานรากตามแนวสายทาง</p>
    <p>พัฒนาโดย รศ.ดร.อิทธิพล มีผล // ภาควิชาครุศาสตร์โยธา // คณะครุศาสตร์อุตสาหกรรม // มจพ.</p>
</div>
""", unsafe_allow_html=True)
